import sys
import os
import re
from docx import Document

def inspect_docx(file_path):
    print(f"Inspecting: {file_path}")
    if not os.path.exists(file_path):
        print("File not found!")
        return

    try:
        doc = Document(file_path)
        
        placeholders = set()
        
        def check_text(text, source):
            matches = re.findall(r'\{\{[^}]+\}\}', text)
            for match in matches:
                placeholders.add(match)
                print(f"Found in {source}: {match}")

        print("\n--- Paragraphs ---")
        for i, para in enumerate(doc.paragraphs):
            if '{{' in para.text:
                check_text(para.text, f"Paragraph {i}")

        print("\n--- Tables ---")
        for i, table in enumerate(doc.tables):
            for r, row in enumerate(table.rows):
                for c, cell in enumerate(row.cells):
                    for p, para in enumerate(cell.paragraphs):
                        if '{{' in para.text:
                            check_text(para.text, f"Table {i} Row {r} Cell {c} Para {p}")

        print("\n--- Headers/Footers ---")
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        if '{{' in para.text:
                            check_text(para.text, "Header")
                    for table in header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    if '{{' in para.text:
                                        check_text(para.text, "Header Table")
            
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        if '{{' in para.text:
                            check_text(para.text, "Footer")
                    for table in footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    if '{{' in para.text:
                                        check_text(para.text, "Footer Table")

        print("\n--- Summary of Placeholders ---")
        for p in sorted(placeholders):
            print(p)

    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    inspect_docx(r"c:\Users\user\Desktop\mygov\my-gov-backend\templates\template_mygov.docx")
