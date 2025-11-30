#!/usr/bin/env python3
"""
Проверка плейсхолдеров в шаблоне
"""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
import re

template_path = '/var/www/mygov-backend/templates/template_mygov.docx'

if not os.path.exists(template_path):
    print(f"✗ Шаблон не найден: {template_path}")
    sys.exit(1)

print("==========================================")
print("  Проверка плейсхолдеров в шаблоне")
print("==========================================")
print(f"\nШаблон: {template_path}")

doc = Document(template_path)

# Паттерн для поиска плейсхолдеров
placeholder_pattern = r'\{\{[^}]+\}\}'

found_placeholders = set()

print("\n1. Поиск плейсхолдеров в параграфах:")
print("-" * 50)
para_count = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text
    placeholders = re.findall(placeholder_pattern, text)
    if placeholders:
        para_count += 1
        print(f"Параграф {i}: {text[:100]}")
        for ph in placeholders:
            found_placeholders.add(ph)
            print(f"  → Найден: {ph}")

print(f"\nВсего параграфов с плейсхолдерами: {para_count}")

print("\n2. Поиск плейсхолдеров в таблицах:")
print("-" * 50)
table_count = 0
for table_idx, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            for para_idx, para in enumerate(cell.paragraphs):
                text = para.text
                placeholders = re.findall(placeholder_pattern, text)
                if placeholders:
                    table_count += 1
                    print(f"Таблица {table_idx}, строка {row_idx}, ячейка {cell_idx}, параграф {para_idx}:")
                    print(f"  Текст: {text[:100]}")
                    for ph in placeholders:
                        found_placeholders.add(ph)
                        print(f"  → Найден: {ph}")

print(f"\nВсего ячеек таблиц с плейсхолдерами: {table_count}")

print("\n3. Поиск плейсхолдеров в колонтитулах:")
print("-" * 50)
header_count = 0
for section in doc.sections:
    # Headers
    for header in [section.header, section.first_page_header, section.even_page_header]:
        if header:
            for para in header.paragraphs:
                text = para.text
                placeholders = re.findall(placeholder_pattern, text)
                if placeholders:
                    header_count += 1
                    print(f"Header: {text[:100]}")
                    for ph in placeholders:
                        found_placeholders.add(ph)
                        print(f"  → Найден: {ph}")
    
    # Footers
    for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
        if footer:
            for para in footer.paragraphs:
                text = para.text
                placeholders = re.findall(placeholder_pattern, text)
                if placeholders:
                    header_count += 1
                    print(f"Footer: {text[:100]}")
                    for ph in placeholders:
                        found_placeholders.add(ph)
                        print(f"  → Найден: {ph}")

print(f"\nВсего колонтитулов с плейсхолдерами: {header_count}")

print("\n4. Все найденные плейсхолдеры:")
print("-" * 50)
if found_placeholders:
    for ph in sorted(found_placeholders):
        print(f"  - {ph}")
else:
    print("  ✗ Плейсхолдеры не найдены!")

print("\n5. Ожидаемые плейсхолдеры (из кода):")
print("-" * 50)
expected = [
    '{{doc_number}}', '{{mygov_doc_number}}', '{{uuid}}',
    '{{patient_name}}', '{{gender}}', '{{age}}', '{{jshshir}}',
    '{{address}}', '{{attached_medical_institution}}',
    '{{diagnosis}}', '{{diagnosis_icd10_code}}',
    '{{final_diagnosis}}', '{{final_diagnosis_icd10_code}}',
    '{{organization}}', '{{doctor_name}}', '{{doctor_position}}',
    '{{department_head_name}}', '{{days_off_from}}', '{{days_off_to}}',
    '{{days_off_period}}', '{{issue_date}}', '{{pin_code}}', '{{qr_code}}'
]

for ph in expected:
    if ph in found_placeholders:
        print(f"  ✓ {ph}")
    else:
        print(f"  ✗ {ph} - НЕ НАЙДЕН в шаблоне!")

print("\n==========================================")
print("  Проверка завершена")
print("==========================================")

