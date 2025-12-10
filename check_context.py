import sys
import os
from docx import Document

def check_context(file_path):
    try:
        doc = Document(file_path)
        
        def print_context(text, source):
            if '{{days_off_period}}' in text:
                print(f"Context in {source}: '{text}'")

        for i, para in enumerate(doc.paragraphs):
            print_context(para.text, f"Paragraph {i}")

        for i, table in enumerate(doc.tables):
            for r, row in enumerate(table.rows):
                for c, cell in enumerate(row.cells):
                    for p, para in enumerate(cell.paragraphs):
                        print_context(para.text, f"Table {i} Row {r} Cell {c} Para {p}")

        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        print_context(para.text, "Header")
                    for table in header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    print_context(para.text, "Header Table")
            
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        print_context(para.text, "Footer")
                    for table in footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    print_context(para.text, "Footer Table")

    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    check_context(r"c:\Users\user\Desktop\mygov\my-gov-backend\templates\template_mygov.docx")
