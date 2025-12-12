"""
MyGov Backend - Генерация документов (DOCX и PDF)
"""
import os
import re
import sys
import uuid
import json
import subprocess
import shutil
import tempfile
from datetime import datetime
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.config import FRONTEND_URL, UPLOAD_FOLDER, TYPE_DOC
from app.services.qr_code import generate_simple_qr, save_qr_to_file
from app.services.storage import storage_manager
from app.services.database import get_next_mygov_doc_number, db_insert, db_update, db_select
from app.utils.logger import (
    logger, log_document_generation, log_pdf_conversion, 
    log_error_with_context, log_function_call
)


@log_function_call
def generate_document(document_data, app=None):
    """
    Генерирует документ MyGov (DOCX и PDF)
    
    Args:
        document_data: Данные документа
        app: Flask приложение (опционально)
    
    Returns:
        dict: Созданный документ или None при ошибке
    """
    document_id = None
    try:
        log_document_generation("START", "Начало генерации документа", 
                               keys=list(document_data.keys()),
                               created_by=document_data.get('created_by'))
        
        # Генерируем уникальные идентификаторы
        document_uuid = str(uuid.uuid4())
        pin_code = generate_pin_code()
        log_document_generation("UUID_GEN", "Сгенерированы идентификаторы", 
                               uuid=document_uuid, pin_code=pin_code)
        
        try:
            mygov_doc_number = get_next_mygov_doc_number()
            log_document_generation("DOC_NUMBER", "Получен номер документа", 
                                   doc_number=mygov_doc_number)
        except Exception as e:
            log_error_with_context(e, f"get_next_mygov_doc_number failed")
            raise
        
        # Стандартный doc_number для совместимости с БД
        doc_number = f"№ MG {mygov_doc_number}"
        
        # Подготавливаем данные для вставки в БД
        db_data = {
            'doc_number': doc_number,
            'pin_code': pin_code,
            'uuid': document_uuid,
            'patient_name': document_data.get('patient_name'),
            'gender': document_data.get('gender'),
            'age': document_data.get('age'),
            'jshshir': document_data.get('jshshir'),
            'address': document_data.get('address'),
            'attached_medical_institution': document_data.get('attached_medical_institution'),
            'diagnosis': document_data.get('diagnosis'),
            'diagnosis_icd10_code': document_data.get('diagnosis_icd10_code'),
            'final_diagnosis': document_data.get('final_diagnosis'),
            'final_diagnosis_icd10_code': document_data.get('final_diagnosis_icd10_code'),
            'organization': document_data.get('organization'),
            'doctor_name': document_data.get('doctor_name'),
            'doctor_position': document_data.get('doctor_position'),
            'department_head_name': document_data.get('department_head_name'),
            'days_off_from': document_data.get('days_off_from'),
            'days_off_to': document_data.get('days_off_to'),
            'issue_date': document_data.get('issue_date', datetime.now().isoformat()),
            'type_doc': TYPE_DOC,  # Всегда 2 для MyGov
            'mygov_doc_number': mygov_doc_number,
            'created_by': document_data.get('created_by')
        }
        
        log_document_generation("DB_INSERT", "Вставка документа в БД", 
                               patient=document_data.get('patient_name'))
        # Вставляем в БД
        created_document = db_insert('documents', db_data)
        if not created_document or not isinstance(created_document, dict):
            logger.error("[DOC_GEN:DB_INSERT] Не удалось создать документ в БД")
            return None
        
        document_id = created_document.get('id')  # type: ignore
        if not document_id:
            logger.error("[DOC_GEN:DB_INSERT] Документ создан, но ID отсутствует")
            return None
        log_document_generation("DB_SUCCESS", "Документ создан в БД", 
                               document_id=document_id, doc_number=mygov_doc_number)
        
        # Генерируем DOCX
        log_document_generation("DOCX_START", "Начало генерации DOCX", document_id=document_id)
        docx_path = fill_docx_template(created_document, app)
        if not docx_path:
            logger.error(f"[DOC_GEN:DOCX_FAIL] Не удалось создать DOCX для документа {document_id}")
            # Удаляем документ из БД если не удалось создать DOCX
            if document_id:
                try:
                    from app.services.database import db_query
                    db_query("DELETE FROM documents WHERE id = %s", [document_id])
                    logger.info(f"[DOC_GEN:CLEANUP] Документ {document_id} удален из БД")
                except Exception as cleanup_error:
                    logger.error(f"[DOC_GEN:CLEANUP] Ошибка при удалении документа: {cleanup_error}")
            return None
        
        log_document_generation("DOCX_SUCCESS", "DOCX создан успешно", 
                               document_id=document_id, docx_path=docx_path)
        
        # Обновляем путь к DOCX в БД
        db_update('documents', {'docx_path': docx_path}, 'id = %s', [document_id])
        log_document_generation("DOCX_UPDATE", "Путь DOCX обновлен в БД", document_id=document_id)
        
        # Конвертируем в PDF
        log_document_generation("PDF_START", "Начало конвертации в PDF", 
                               document_id=document_id, docx_path=docx_path)
        pdf_path = convert_docx_to_pdf(docx_path, document_uuid, app)
        if pdf_path:
            log_document_generation("PDF_SUCCESS", "PDF создан успешно", 
                                   document_id=document_id, pdf_path=pdf_path)
            try:
                db_update('documents', {'pdf_path': pdf_path}, 'id = %s', [document_id])
                log_document_generation("PDF_UPDATE", "Путь PDF обновлен в БД", document_id=document_id, pdf_path=pdf_path)
                
                # Проверяем, что обновление прошло успешно
                updated_doc = db_select('documents', 'id = %s', [document_id], fetch_one=True)
                if updated_doc and isinstance(updated_doc, dict) and updated_doc.get('pdf_path') == pdf_path:
                    logger.debug(f"[DOC_GEN:PDF_VERIFY] PDF путь подтвержден в БД: {pdf_path}")
                else:
                    pdf_path_actual = updated_doc.get('pdf_path') if (updated_doc and isinstance(updated_doc, dict)) else 'None'
                    logger.error(f"[DOC_GEN:PDF_VERIFY_FAIL] PDF путь не обновлен в БД! Ожидалось: {pdf_path}, получено: {pdf_path_actual}")
            except Exception as update_error:
                log_error_with_context(update_error, f"Ошибка при обновлении pdf_path в БД, document_id={document_id}, pdf_path={pdf_path}")
        else:
            logger.warning(f"[DOC_GEN:PDF_WARNING] PDF не был создан для документа {document_id}, но документ сохранен")
        
        # Возвращаем результат
        if isinstance(created_document, dict):
            created_document['docx_path'] = docx_path  # type: ignore
            created_document['pdf_path'] = pdf_path  # type: ignore
        
        log_document_generation("SUCCESS", "Генерация документа завершена успешно", 
                               document_id=document_id, doc_number=mygov_doc_number,
                               has_docx=bool(docx_path), has_pdf=bool(pdf_path))
        return created_document
        
    except Exception as e:
        log_error_with_context(e, f"generate_document failed, document_id={document_id}")
        
        # Удаляем документ из БД если была ошибка после создания
        if document_id:
            try:
                from app.services.database import db_query
                logger.info(f"[DOC_GEN:CLEANUP] Удаление документа {document_id} из БД из-за ошибки")
                db_query("DELETE FROM documents WHERE id = %s", [document_id])
            except Exception as cleanup_error:
                logger.error(f"[DOC_GEN:CLEANUP] Ошибка при очистке: {cleanup_error}")
        
        return None


def generate_pin_code():
    """Генерирует 4-значный PIN-код"""
    import random
    return ''.join([str(random.randint(0, 9)) for _ in range(4)])


def fill_docx_template(document_data, app=None):
    """
    Заполняет шаблон DOCX данными документа
    
    Args:
        document_data: Данные документа
        app: Flask приложение
    
    Returns:
        str: Путь к созданному файлу или None
    """
    try:
        template_folder = 'templates'
        if app:
            template_folder = app.config.get('TEMPLATE_FOLDER', 'templates')
        
        # Получаем абсолютный путь к шаблону
        # app.root_path указывает на директорию приложения (app/), нужно подняться на уровень выше
        if app:
            # app.root_path = /var/www/mygov-backend/app
            # Нужно получить /var/www/mygov-backend
            project_root = os.path.dirname(app.root_path)
            template_path = os.path.join(project_root, template_folder, 'template_mygov.docx')
            logger.debug(f"[DOCX_TEMPLATE] app.root_path: {app.root_path}, project_root: {project_root}")
        else:
            # Используем относительный путь от текущей директории
            template_path = os.path.join(template_folder, 'template_mygov.docx')
            # Пробуем абсолютный путь
            if not os.path.exists(template_path):
                # Пробуем от корня проекта (3 уровня вверх от app/services/document.py)
                current_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
                template_path = os.path.join(current_dir, template_folder, 'template_mygov.docx')
        
        logger.debug(f"[DOCX_TEMPLATE] Поиск шаблона: {template_path}")
        
        # Дополнительные попытки найти шаблон
        if not os.path.exists(template_path):
            # Попробуем стандартные пути
            alternative_paths = [
                os.path.join('/var/www/mygov-backend', template_folder, 'template_mygov.docx'),
                os.path.join('/var/www/mygov-backend', 'templates', 'template_mygov.docx'),
                os.path.join(os.getcwd(), template_folder, 'template_mygov.docx'),
                os.path.join(os.getcwd(), 'templates', 'template_mygov.docx'),
            ]
            
            for alt_path in alternative_paths:
                if os.path.exists(alt_path):
                    logger.info(f"[DOCX_TEMPLATE:FOUND_ALT] Шаблон найден по альтернативному пути: {alt_path}")
                    template_path = alt_path
                    break
            else:
                logger.error(f"[DOCX_TEMPLATE:NOT_FOUND] Шаблон не найден: {template_path}")
                logger.debug(f"[DOCX_TEMPLATE] Текущая рабочая директория: {os.getcwd()}")
                logger.debug(f"[DOCX_TEMPLATE] Абсолютный путь к скрипту: {os.path.abspath(__file__)}")
                if app:
                    logger.debug(f"[DOCX_TEMPLATE] app.root_path: {app.root_path}")
                    logger.debug(f"[DOCX_TEMPLATE] Проверка существования директории templates: {os.path.exists(os.path.join(os.path.dirname(app.root_path), 'templates'))}")
                return None
        
        logger.debug(f"[DOCX_TEMPLATE:FOUND] Шаблон найден: {template_path}")
        
        doc = Document(template_path)
        logger.debug(f"[DOCX_TEMPLATE] Документ загружен из шаблона")
        
        # Подготавливаем замены
        logger.debug(f"[DOCX_TEMPLATE] Подготовка замен плейсхолдеров...")
        replacements = prepare_replacements(document_data)
        logger.info(f"[DOCX_TEMPLATE] Подготовлено {len(replacements)} замен")
        logger.debug(f"[DOCX_TEMPLATE] Примеры замен: {list(replacements.keys())[:5]}")
        
        # Логируем некоторые значения для проверки
        sample_keys = ['{{patient_name}}', '{{doc_number}}', '{{organization}}', '{{doctor_name}}']
        for key in sample_keys:
            if key in replacements:
                logger.debug(f"[DOCX_TEMPLATE] {key} = '{replacements[key]}'")
        
        # Заменяем плейсхолдеры
        logger.debug(f"[DOCX_TEMPLATE] Начало замены плейсхолдеров в документе...")
        replace_placeholders(doc, replacements)
        logger.debug(f"[DOCX_TEMPLATE] Замена плейсхолдеров завершена")
        
        # Форматируем метки полей жирным шрифтом
        logger.debug(f"[DOCX_TEMPLATE] Форматирование меток полей жирным шрифтом...")
        format_field_labels(doc)
        logger.debug(f"[DOCX_TEMPLATE] Форматирование меток полей завершено")
        
        # Добавляем QR-код
        add_qr_code(doc, document_data, app)
        
        # Сохраняем документ
        #region agent log
        try:
            with open(r'c:\Users\user\Desktop\mygov\my-gov-backend\.cursor\debug.log', 'a', encoding='utf-8') as _f:
                _f.write(json.dumps({
                    "sessionId": "debug-session",
                    "runId": "pre-fix",
                    "hypothesisId": "H1",
                    "location": "document.py:fill_docx_template:pre_makedirs",
                    "message": "Preparing to ensure upload folder",
                    "data": {
                        "cwd": os.getcwd(),
                        "upload_folder": UPLOAD_FOLDER,
                        "abs_upload": os.path.abspath(UPLOAD_FOLDER),
                        "exists": os.path.exists(UPLOAD_FOLDER),
                        "writable": os.access(UPLOAD_FOLDER, os.W_OK),
                        "uid": getattr(os, "getuid", lambda: None)(),
                        "euid": getattr(os, "geteuid", lambda: None)(),
                        "gid": getattr(os, "getgid", lambda: None)(),
                    },
                    "timestamp": int(datetime.now().timestamp() * 1000)
                }) + "\n")
        except Exception:
            pass
        #endregion
        os.makedirs(UPLOAD_FOLDER, exist_ok=True)
        #region agent log
        try:
            with open(r'c:\Users\user\Desktop\mygov\my-gov-backend\.cursor\debug.log', 'a', encoding='utf-8') as _f:
                _f.write(json.dumps({
                    "sessionId": "debug-session",
                    "runId": "pre-fix",
                    "hypothesisId": "H1",
                    "location": "document.py:fill_docx_template:post_makedirs",
                    "message": "Upload folder ensured",
                    "data": {
                        "abs_upload": os.path.abspath(UPLOAD_FOLDER),
                        "exists": os.path.exists(UPLOAD_FOLDER),
                        "writable": os.access(UPLOAD_FOLDER, os.W_OK),
                        "stat_mode": oct(os.stat(UPLOAD_FOLDER).st_mode) if os.path.exists(UPLOAD_FOLDER) else None,
                        "uid": getattr(os, "getuid", lambda: None)(),
                        "gid": getattr(os, "getgid", lambda: None)(),
                    },
                    "timestamp": int(datetime.now().timestamp() * 1000)
                }) + "\n")
        except Exception:
            pass
        #endregion
        document_uuid = document_data.get('uuid', str(uuid.uuid4()))
        output_path = os.path.join(UPLOAD_FOLDER, f"{document_uuid}.docx")
        
        doc.save(output_path)
        #region agent log
        try:
            with open(r'c:\Users\user\Desktop\mygov\my-gov-backend\.cursor\debug.log', 'a', encoding='utf-8') as _f:
                _f.write(json.dumps({
                    "sessionId": "debug-session",
                    "runId": "pre-fix",
                    "hypothesisId": "H1",
                    "location": "document.py:fill_docx_template:post_doc_save",
                    "message": "DOCX saved locally",
                    "data": {
                        "output_path": output_path,
                        "exists": os.path.exists(output_path),
                        "writable": os.access(os.path.dirname(output_path) or ".", os.W_OK),
                        "size": os.path.getsize(output_path) if os.path.exists(output_path) else None
                    },
                    "timestamp": int(datetime.now().timestamp() * 1000)
                }) + "\n")
        except Exception:
            pass
        #endregion
        
        # Сохраняем в MinIO
        with open(output_path, 'rb') as f:
            docx_data = f.read()
        
        stored_path = storage_manager.save_file(
            docx_data, 
            f"{document_uuid}.docx",
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Удаляем локальный файл если сохранили в MinIO
        if storage_manager.use_minio and os.path.exists(output_path):
            try:
                os.remove(output_path)
            except:
                pass
        
        return stored_path
        
    except Exception as e:
        log_error_with_context(e, "fill_docx_template failed")
        return None


def prepare_replacements(document_data):
    """Подготавливает словарь замен для шаблона"""
    # Форматируем даты в формат DD.MM.YYYY
    def format_date(date_str):
        if not date_str:
            return ''
        try:
            # Преобразуем в строку
            date_str = str(date_str).strip()
            
            # Если есть время (ISO формат), берем только дату
            if 'T' in date_str:
                date_str = date_str.split('T')[0]
            # Если есть пробел, берем только дату
            elif ' ' in date_str:
                date_str = date_str.split(' ')[0]
            
            # Парсим дату в формате YYYY-MM-DD
            dt = datetime.strptime(date_str, '%Y-%m-%d')
            # Возвращаем в формате DD.MM.YYYY
            return dt.strftime('%d.%m.%Y')
        except Exception as e:
            # Если не удалось распарсить, возвращаем исходную строку
            logger.warning(f"[DATE_FORMAT] Не удалось отформатировать дату '{date_str}': {e}")
            return str(date_str) if date_str else ''
            
    # Вычисляем количество дней
    def calculate_days_off(start_date, end_date):
        try:
            if not start_date or not end_date:
                return ''
            if 'T' in str(start_date): start_date = str(start_date).split('T')[0]
            if 'T' in str(end_date): end_date = str(end_date).split('T')[0]
            
            d1 = datetime.strptime(str(start_date), '%Y-%m-%d')
            d2 = datetime.strptime(str(end_date), '%Y-%m-%d')
            delta = (d2 - d1).days + 1 # Включая первый день
            return str(delta)
        except:
            return ''
    
    # days_off_count = calculate_days_off(document_data.get('days_off_from'), document_data.get('days_off_to'))
    
    # Формируем период освобождения (даты)
    date_from = format_date(document_data.get('days_off_from'))
    date_to = format_date(document_data.get('days_off_to'))
    
    days_off_period_str = ""
    if date_from and date_to:
        days_off_period_str = f"{date_from} - {date_to}"
    elif date_from:
        days_off_period_str = f"с {date_from}"
    
    return {
        '{{doc_number}}': f"№ {document_data.get('mygov_doc_number', '')}", # Форматированный номер
        '{{mygov_doc_number}}': document_data.get('mygov_doc_number', ''),
        '{{uuid}}': document_data.get('uuid', ''),
        '{{patient_name}}': document_data.get('patient_name', ''),
        '{{gender}}': document_data.get('gender', ''),
        '{{age}}': str(document_data.get('age', '')),
        '{{jshshir}}': document_data.get('jshshir', ''),
        '{{address}}': document_data.get('address', ''),
        '{{attached_medical_institution}}': document_data.get('attached_medical_institution', ''),
        '{{diagnosis}}': document_data.get('diagnosis', ''),
        '{{diagnosis_icd10_code}}': document_data.get('diagnosis_icd10_code', ''),
        '{{final_diagnosis}}': document_data.get('final_diagnosis', ''),
        '{{final_diagnosis_icd10_code}}': document_data.get('final_diagnosis_icd10_code', ''),
        '{{organization}}': document_data.get('organization', ''),
        '{{doctor_name}}': document_data.get('doctor_name', ''),
        '{{doctor_position}}': document_data.get('doctor_position', ''),
        '{{department_head_name}}': document_data.get('department_head_name', ''),
        '{{days_off_from}}': date_from,
        '{{days_off_to}}': date_to,
        '{{days_off_period}}': days_off_period_str,
        '{{issue_date}}': format_date(document_data.get('issue_date')),
        '{{pin_code}}': document_data.get('pin_code', ''),  # PIN-код заменяется здесь
        # {{qr_code}} обрабатывается отдельно в add_qr_code (нужно изображение)
    }


def replace_placeholders(doc, replacements):
    """Заменяет плейсхолдеры в документе"""
    total_replacements = 0
    found_placeholders = []
    
    def process_paragraph(paragraph):
        """Обрабатывает один параграф"""
        nonlocal total_replacements
        text = paragraph.text
        
        # Пропускаем параграфы с {{qr_code}} - он обрабатывается отдельно (нужно изображение)
        if '{{qr_code}}' in text:
            return
        
        has_replacement = False
        replaced_keys = []
        for key, value in replacements.items():
            if key in text:
                text = text.replace(key, str(value))
                has_replacement = True
                replaced_keys.append(key)
                total_replacements += 1
                if key not in found_placeholders:
                    found_placeholders.append(key)
        
        if has_replacement:
            # Если была замена, обновляем текст параграфа
            # Сохраняем стиль исходного параграфа (первого run)
            style = None
            if paragraph.runs:
                style = paragraph.runs[0].style
                
            paragraph.text = text
            
            # Устанавливаем шрифт Times New Roman 10.5pt для всех runs
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                
                # Проверяем, является ли текст PIN-кодом (4 цифры)
                is_pin_code = run.text.strip().isdigit() and len(run.text.strip()) == 4
                
                if is_pin_code:
                    # PIN-код: размер 23pt, жирный
                    run.font.size = Pt(18)
                    run.font.bold = True
                else:
                    # Обычный текст: размер 10.5pt
                    run.font.size = Pt(10.5)
                
                # Устанавливаем шрифт для кириллицы (важно для русского и узбекского текста)
                # Проверяем наличие rPr элемента
                if run._element.rPr is None:
                    run._element.get_or_add_rPr()
                
                # Устанавливаем rFonts для правильного отображения кириллицы
                rFonts = run._element.rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement(qn('w:rFonts'))
                    run._element.rPr.append(rFonts)
                
                rFonts.set(qn('w:ascii'), 'Times New Roman')
                rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                rFonts.set(qn('w:cs'), 'Times New Roman')
                
                # Восстанавливаем стиль если был
                if style:
                    run.style = style

    logger.debug(f"[DOCX_REPLACE] Начало обработки документа: {len(doc.paragraphs)} параграфов, {len(doc.tables)} таблиц")
    
    # 1. Обрабатываем тело документа
    for para in doc.paragraphs:
        process_paragraph(para)
        
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                    for para in cell.paragraphs:
                        process_paragraph(para)
                        
    # 2. Обрабатываем колонтитулы (Headers & Footers)
    logger.debug(f"[DOCX_REPLACE] Обработка колонтитулов...")
    for section in doc.sections:
        # Headers
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for para in header.paragraphs:
                    process_paragraph(para)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                process_paragraph(para)
                                
        # Footers
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    process_paragraph(para)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                process_paragraph(para)
    
    logger.info(f"[DOCX_REPLACE] Замена завершена: всего замен {total_replacements}, найдено плейсхолдеров: {len(found_placeholders)}")
    if found_placeholders:
        logger.debug(f"[DOCX_REPLACE] Найденные плейсхолдеры: {found_placeholders}")
    else:
        logger.warning(f"[DOCX_REPLACE] ВНИМАНИЕ: Плейсхолдеры не найдены в документе! Проверьте шаблон.")


def format_field_labels(doc):
    """
    Форматирует все метки полей (названия переменных) и тексты перед значениями жирным шрифтом.
    Автоматически находит тексты, заканчивающиеся двоеточием, и делает их жирными.
    """
    
    def process_paragraph(paragraph):
        """Обрабатывает параграф и делает все метки полей (тексты с двоеточием) жирными"""
        if not paragraph.runs:
            return
        
        # Получаем весь текст параграфа
        full_text = paragraph.text
        
        # Ищем все тексты, заканчивающиеся двоеточием (метки полей)
        # Паттерн: текст, заканчивающийся двоеточием, возможно с пробелами после
        # Это покроет все метки типа "F.I.Sh:", "Jinsi:", "Yashash manzili:" и т.д.
        pattern = r'([^\s:]+(?:\s+[^\s:]+)*\s*:)'
        
        # Проверяем, есть ли в тексте метки полей
        matches = list(re.finditer(pattern, full_text))
        if not matches:
            return
        
        # Если нашли метки, нужно переформатировать параграф
        # Сохраняем выравнивание
        alignment = paragraph.alignment
        
        # Очищаем параграф
        paragraph.clear()
        if alignment:
            paragraph.alignment = alignment
        
        # Разбиваем текст на части по меткам полей
        last_end = 0
        for match in matches:
            # Добавляем текст до метки
            if match.start() > last_end:
                text_before = full_text[last_end:match.start()]
                if text_before:
                    run = paragraph.add_run(text_before)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10.5)
                    if run._element.rPr is None:
                        run._element.get_or_add_rPr()
                    rFonts = run._element.rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement(qn('w:rFonts'))
                        run._element.rPr.append(rFonts)
                    rFonts.set(qn('w:ascii'), 'Times New Roman')
                    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    rFonts.set(qn('w:cs'), 'Times New Roman')
            
            # Добавляем метку поля жирным
            label_text = match.group(1)
            run = paragraph.add_run(label_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            run.font.bold = True  # Делаем жирным
            logger.debug(f"[DOCX_FORMAT] Метка поля выделена жирным: {label_text}")
            
            if run._element.rPr is None:
                run._element.get_or_add_rPr()
            rFonts = run._element.rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement(qn('w:rFonts'))
                run._element.rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rFonts.set(qn('w:cs'), 'Times New Roman')
            
            last_end = match.end()
        
        # Добавляем оставшийся текст после последней метки
        if last_end < len(full_text):
            text_after = full_text[last_end:]
            if text_after:
                run = paragraph.add_run(text_after)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10.5)
                if run._element.rPr is None:
                    run._element.get_or_add_rPr()
                rFonts = run._element.rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement(qn('w:rFonts'))
                    run._element.rPr.append(rFonts)
                rFonts.set(qn('w:ascii'), 'Times New Roman')
                rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                rFonts.set(qn('w:cs'), 'Times New Roman')
    
    # Обрабатываем все параграфы в документе
    for para in doc.paragraphs:
        process_paragraph(para)
    
    # Обрабатываем параграфы в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para)
    
    # Обрабатываем колонтитулы
    for section in doc.sections:
        # Headers
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for para in header.paragraphs:
                    process_paragraph(para)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                process_paragraph(para)
        
        # Footers
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    process_paragraph(para)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                process_paragraph(para)
    
    logger.debug(f"[DOCX_FORMAT] Форматирование меток полей завершено")


def add_qr_code(doc, document_data, app=None):
    """Добавляет QR-код и PIN-код в документ inline в том же параграфе"""
    try:
        document_uuid = document_data.get('uuid', '')
        pin_code = document_data.get('pin_code', '')
        
        # Генерируем URL для QR-кода (ссылка на страницу проверки документа)
        frontend_url = FRONTEND_URL.rstrip('/')
        qr_url = f"{frontend_url}/verify/{document_uuid}"
        
        logger.info(f"[QR_CODE] Генерация QR-кода для проверки документа: {qr_url}")
        logger.debug(f"[QR_CODE] UUID: {document_uuid}, PIN: {pin_code}")
        
        # Генерируем простой QR-код
        qr_img = generate_simple_qr(qr_url, box_size=3, border=1)
        
        # Сохраняем во временный файл
        os.makedirs(UPLOAD_FOLDER, exist_ok=True)
        qr_temp_path = os.path.join(UPLOAD_FOLDER, f'qr_temp_{pin_code}.png')
        #region agent log
        try:
            with open(r'c:\Users\user\Desktop\mygov\my-gov-backend\.cursor\debug.log', 'a', encoding='utf-8') as _f:
                _f.write(json.dumps({
                    "sessionId": "debug-session",
                    "runId": "pre-fix",
                    "hypothesisId": "H2",
                    "location": "document.py:add_qr_code:pre_save_qr",
                    "message": "About to save QR to temp file",
                    "data": {
                        "qr_temp_path": qr_temp_path,
                        "upload_exists": os.path.exists(UPLOAD_FOLDER),
                        "upload_writable": os.access(UPLOAD_FOLDER, os.W_OK),
                        "temp_exists": os.path.exists(qr_temp_path),
                    },
                    "timestamp": int(datetime.now().timestamp() * 1000)
                }) + "\n")
        except Exception:
            pass
        #endregion
        save_qr_to_file(qr_img, qr_temp_path)
        #region agent log
        try:
            with open(r'c:\Users\user\Desktop\mygov\my-gov-backend\.cursor\debug.log', 'a', encoding='utf-8') as _f:
                _f.write(json.dumps({
                    "sessionId": "debug-session",
                    "runId": "pre-fix",
                    "hypothesisId": "H2",
                    "location": "document.py:add_qr_code:post_save_qr",
                    "message": "QR temp file saved",
                    "data": {
                        "qr_temp_path": qr_temp_path,
                        "exists": os.path.exists(qr_temp_path),
                        "size": os.path.getsize(qr_temp_path) if os.path.exists(qr_temp_path) else None
                    },
                    "timestamp": int(datetime.now().timestamp() * 1000)
                }) + "\n")
        except Exception:
            pass
        #endregion
        
        qr_added = False
        
        def replace_pin_qr_inline(paragraph):
            """Просто заменяет {{pin_code}} и {{qr_code}} inline, сохраняя структуру шаблона"""
            nonlocal qr_added
            text = paragraph.text
            
            # Проверяем наличие плейсхолдеров
            has_qr = '{{qr_code}}' in text
            
            if not has_qr:
                return False
            
            # Сохраняем выравнивание параграфа
            alignment = paragraph.alignment
            
            # Сохраняем полный текст для разбора
            full_text = text
            
            # Очищаем параграф
            paragraph.clear()
            if alignment:
                paragraph.alignment = alignment
            
            # Разбиваем текст на части по плейсхолдерам
            parts = re.split(r'({{pin_code}}|{{qr_code}})', full_text)
            
            for part in parts:
                if not part:
                    continue
                
                if part == '{{pin_code}}':
                    # Добавляем PIN-код как текст
                    run = paragraph.add_run(pin_code)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(18)
                    run.font.bold = True
                    
                    # Устанавливаем шрифт для кириллицы
                    if run._element.rPr is None:
                        run._element.get_or_add_rPr()
                    rFonts = run._element.rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement(qn('w:rFonts'))
                        run._element.rPr.append(rFonts)
                    rFonts.set(qn('w:ascii'), 'Times New Roman')
                    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    rFonts.set(qn('w:cs'), 'Times New Roman')
                    
                elif part == '{{qr_code}}':
                    # Добавляем QR-код как изображение
                    run = paragraph.add_run()
                    run.add_picture(qr_temp_path, width=Inches(0.8))
                    
                else:
                    # Обычный текст - сохраняем как есть
                    run = paragraph.add_run(part)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10.5)
                    
                    # Устанавливаем шрифт для кириллицы
                    if run._element.rPr is None:
                        run._element.get_or_add_rPr()
                    rFonts = run._element.rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement(qn('w:rFonts'))
                        run._element.rPr.append(rFonts)
                    rFonts.set(qn('w:ascii'), 'Times New Roman')
                    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    rFonts.set(qn('w:cs'), 'Times New Roman')
            
            return True
        
        # Вспомогательная функция для обработки параграфов
        def check_paragraphs(paragraphs):
            nonlocal qr_added
            for para in paragraphs:
                if replace_pin_qr_inline(para):
                    qr_added = True
                    return True
            return False
            
        # Вспомогательная функция для обработки таблиц
        def check_tables(tables):
            nonlocal qr_added
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        if check_paragraphs(cell.paragraphs):
                            return True
            return False

        # 1. Проверяем тело документа
        if not check_paragraphs(doc.paragraphs):
            if not check_tables(doc.tables):
                # 2. Проверяем Headers и Footers
                for section in doc.sections:
                    # Footers (чаще всего QR там)
                    for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                        if footer:
                            if check_paragraphs(footer.paragraphs): break
                            if check_tables(footer.tables): break
                    if qr_added: break
                    
                    # Headers (на всякий случай)
                    for header in [section.header, section.first_page_header, section.even_page_header]:
                        if header:
                            if check_paragraphs(header.paragraphs): break
                            if check_tables(header.tables): break
                    if qr_added: break
        
        # Если плейсхолдер не найден, добавляем в конец документа
        if not qr_added:
            add_pin_qr_to_end(doc, pin_code, qr_temp_path)
        
        # Удаляем временный файл
        if os.path.exists(qr_temp_path):
            os.remove(qr_temp_path)
            
    except Exception as e:
        log_error_with_context(e, "add_qr_code failed")


def add_pin_qr_table(doc, para, pin_code, qr_temp_path):
    """Добавляет таблицу с PIN-кодом и QR-кодом"""
    parent = para._element.getparent()
    
    # Создаем таблицу
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Cm(2.0)
    table.columns[1].width = Cm(2.5)
    
    # Ячейка с PIN-кодом
    pin_cell = table.rows[0].cells[0]
    pin_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    pin_para = pin_cell.paragraphs[0]
    pin_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pin_run = pin_para.add_run(pin_code)
    pin_run.font.size = Pt(18)
    pin_run.font.bold = True
    
    # Ячейка с QR-кодом
    qr_cell = table.rows[0].cells[1]
    qr_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    qr_para = qr_cell.paragraphs[0]
    qr_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    qr_run = qr_para.add_run()
    qr_run.add_picture(qr_temp_path, width=Inches(0.8))
    
    # Вставляем таблицу вместо параграфа
    parent.insert(parent.index(para._element), table._element)
    parent.remove(para._element)


def add_pin_qr_to_end(doc, pin_code, qr_temp_path):
    """Добавляет PIN-код и QR-код в конец документа"""
    # Добавляем таблицу
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Cm(2.0)
    table.columns[1].width = Cm(2.5)
    
    # PIN-код
    pin_cell = table.rows[0].cells[0]
    pin_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    pin_para = pin_cell.paragraphs[0]
    pin_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pin_run = pin_para.add_run(pin_code)
    pin_run.font.size = Pt(23)
    pin_run.font.bold = True
    
    # QR-код
    qr_cell = table.rows[0].cells[1]
    qr_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    qr_para = qr_cell.paragraphs[0]
    qr_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    qr_run = qr_para.add_run()
    qr_run.add_picture(qr_temp_path, width=Inches(0.8))


@log_function_call
def convert_docx_to_pdf(docx_path, document_uuid, app=None):
    """
    Конвертирует DOCX в PDF с помощью LibreOffice
    
    Args:
        docx_path: Путь к DOCX файлу
        document_uuid: UUID документа
        app: Flask приложение
    
    Returns:
        str: Путь к PDF файлу или None
    """
    temp_docx = None
    try:
        log_pdf_conversion("START", "Начало конвертации DOCX в PDF", 
                          docx_path=docx_path, uuid=document_uuid)
        
        # Создаем уникальную временную директорию для этого процесса
        # Это решает проблему прав доступа, если /tmp/.config создан рутом или другим пользователем
        with tempfile.TemporaryDirectory(prefix=f'lo_conv_{document_uuid}_') as unique_temp_dir:
            # Используем уникальную директорию как output_dir и HOME
            output_dir = unique_temp_dir
            log_pdf_conversion("TEMP_DIR", "Создана уникальная временная директория", dir=output_dir)

            # Получаем DOCX из хранилища
            if docx_path.startswith('minio://'):
                log_pdf_conversion("MINIO_GET", "Получение DOCX из MinIO", minio_path=docx_path)
                docx_data = storage_manager.get_file(docx_path)
                if not docx_data:
                    logger.error(f"[PDF_CONV:MINIO_FAIL] Не удалось получить DOCX из MinIO: {docx_path}")
                    return None
                
                # Сохраняем во временный файл в нашей уникальной директории
                temp_docx = os.path.join(output_dir, f"input_{document_uuid}.docx")
                log_pdf_conversion("TEMP_SAVE", "Сохранение временного DOCX", temp_path=temp_docx)
                with open(temp_docx, 'wb') as f:
                    f.write(docx_data)
                # Устанавливаем права
                os.chmod(temp_docx, 0o666)
                docx_path = temp_docx
                
            elif not os.path.exists(docx_path):
                # Старый формат пути (локальный файл) - пробуем получить из MinIO по UUID
                logger.warning(f"[PDF_CONV:OLD_FORMAT] Обнаружен старый формат пути, пробуем получить из MinIO: {docx_path}")
                
                # Пробуем несколько вариантов получения файла из MinIO
                possible_paths = []
                
                # Вариант 1: Используем UUID напрямую
                if document_uuid:
                    possible_paths.append(f"minio://dmed/{document_uuid}.docx")
                
                # Вариант 2: Извлекаем имя файла из старого пути
                filename = os.path.basename(docx_path)
                if filename and filename.endswith('.docx'):
                    possible_paths.append(f"minio://dmed/{filename}")
                    # Вариант 3: Без расширения
                    name_without_ext = filename.replace('.docx', '')
                    if name_without_ext:
                        possible_paths.append(f"minio://dmed/{name_without_ext}.docx")
                
                docx_data = None
                used_path = None
                for minio_path in possible_paths:
                    log_pdf_conversion("MINIO_GET_OLD", "Попытка получить DOCX из MinIO", minio_path=minio_path)
                    docx_data = storage_manager.get_file(minio_path)
                    if docx_data:
                        used_path = minio_path
                        logger.info(f"[PDF_CONV:MINIO_GET_OLD_SUCCESS] DOCX получен из MinIO: {minio_path}")
                        break
                
                if docx_data:
                    # Сохраняем во временный файл
                    temp_docx = os.path.join(output_dir, f"input_{document_uuid}.docx")
                    with open(temp_docx, 'wb') as f:
                        f.write(docx_data)
                    os.chmod(temp_docx, 0o666)
                    docx_path = temp_docx
                    log_pdf_conversion("TEMP_SAVE_OLD", "Сохранение DOCX из старого формата", temp_path=temp_docx, minio_path=used_path)
                else:
                    logger.error(f"[PDF_CONV:OLD_FORMAT_FAIL] Не удалось получить DOCX из MinIO. Пробовали пути: {possible_paths}")
                    return None
            
            # Проверяем существование файла
            if not os.path.exists(docx_path):
                logger.error(f"[PDF_CONV:FILE_NOT_FOUND] DOCX файл не найден: {docx_path}")
                return None
            
            file_size = os.path.getsize(docx_path)
            log_pdf_conversion("FILE_CHECK", "DOCX файл проверен", 
                              docx_path=docx_path, file_size=file_size)
            
            # Конвертируем с помощью LibreOffice
            libreoffice_cmd = find_libreoffice()
            if not libreoffice_cmd:
                logger.warning("[PDF_CONV:LIBREOFFICE_NOT_FOUND] LibreOffice не найден, PDF не будет создан")
                return None
            
            log_pdf_conversion("LIBREOFFICE_FOUND", "LibreOffice найден", cmd=libreoffice_cmd)
            
            # Абсолютный путь к входному файлу
            abs_docx_path = os.path.abspath(docx_path)
            log_pdf_conversion("ABS_PATH", "Абсолютный путь к DOCX", abs_path=abs_docx_path)
            
            # Формируем команду конвертации
            # Используем минимальный набор флагов для совместимости (как в dmed)
            cmd = [
                libreoffice_cmd,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', output_dir,
                abs_docx_path
            ]
            
            # Для Windows добавляем специальную директорию конфигурации (как в dmed)
            temp_config_dir = None
            if sys.platform == 'win32':
                temp_config_dir = tempfile.mkdtemp(prefix='LibreOffice_Config_')
                config_path = os.path.abspath(temp_config_dir).replace('\\', '/')
                if not config_path.startswith('/'):
                    config_path = '/' + config_path
                # Используем правильный формат URI для Windows
                cmd.extend(['-env:UserInstallation=file:///' + config_path.lstrip('/')])
                log_pdf_conversion("WINDOWS_CONFIG", "Используется отдельная директория конфигурации для Windows", 
                                  config_dir=temp_config_dir)
            
            log_pdf_conversion("CMD", "Команда конвертации", cmd=' '.join(cmd))
            
            # Устанавливаем переменные окружения для LibreOffice
            env = os.environ.copy()
            
            # КРИТИЧНО: Устанавливаем полный PATH для Linux
            # LibreOffice - это shell-скрипт, который использует dirname, basename, ls, sed, grep, uname
            # Если PATH не установлен правильно, эти команды не найдутся (код ошибки 127)
            if sys.platform != 'win32':
                # Устанавливаем стандартный Linux PATH
                env['PATH'] = '/usr/local/sbin:/usr/local/bin:/usr/sbin:/usr/bin:/sbin:/bin'
                # Используем output_dir как HOME, чтобы LibreOffice мог писать в .config
                env['HOME'] = output_dir
            
            env['TMPDIR'] = output_dir
            env['TMP'] = output_dir
            env['TEMP'] = output_dir
            
            log_pdf_conversion("EXEC_START", "Запуск LibreOffice", 
                              env_home=env.get('HOME', 'default'), env_tmpdir=env['TMPDIR'],
                              platform=sys.platform)
            
            # Увеличиваем таймаут до 180 секунд (как в dmed) для больших документов
            result = subprocess.run(
                cmd, 
                capture_output=True, 
                text=True, 
                timeout=180,
                env=env,
                cwd=output_dir
            )
            
            log_pdf_conversion("EXEC_RESULT", "LibreOffice завершен", 
                              returncode=result.returncode,
                              stdout_len=len(result.stdout) if result.stdout else 0,
                              stderr_len=len(result.stderr) if result.stderr else 0)
            
            if result.stdout:
                logger.debug(f"[PDF_CONV:STDOUT] {result.stdout[:500]}")
            if result.stderr:
                logger.debug(f"[PDF_CONV:STDERR] {result.stderr[:500]}")
            
            # Улучшенная обработка ошибок (как в dmed)
            if result.returncode != 0:
                logger.error(f"[PDF_CONV:EXEC_FAIL] LibreOffice вернул код {result.returncode}")
                if result.stderr:
                    logger.error(f"[PDF_CONV:STDERR_FULL] {result.stderr}")
                    stderr_lower = result.stderr.lower()
                    # Проверяем конкретные ошибки
                    if 'bootstrap.ini' in stderr_lower and ('повреждён' in stderr_lower or 'damaged' in stderr_lower or 'corrupted' in stderr_lower):
                        logger.error("[PDF_CONV:BOOTSTRAP_ERROR] Файл конфигурации LibreOffice повреждён")
                    elif 'document is empty' in stderr_lower or 'source file could not be loaded' in stderr_lower:
                        logger.error(f"[PDF_CONV:LOAD_ERROR] LibreOffice не может загрузить файл: {abs_docx_path}")
                        # Дополнительная диагностика
                        try:
                            import zipfile
                            with zipfile.ZipFile(abs_docx_path, 'r') as zf:
                                file_list = zf.namelist()
                                logger.debug(f"[PDF_CONV:ZIP_CONTENTS] Содержимое DOCX (первые 10 файлов): {file_list[:10]}")
                                required_files = ['[Content_Types].xml', 'word/document.xml']
                                missing = [f for f in required_files if f not in file_list]
                                if missing:
                                    logger.error(f"[PDF_CONV:MISSING_FILES] В DOCX отсутствуют обязательные файлы: {missing}")
                        except Exception as zip_error:
                            logger.error(f"[PDF_CONV:ZIP_ERROR] Не удалось прочитать DOCX как ZIP: {zip_error}")
                if result.stdout:
                    logger.error(f"[PDF_CONV:STDOUT_FULL] {result.stdout}")
                
                # Очистка временной директории конфигурации для Windows
                if temp_config_dir and os.path.exists(temp_config_dir):
                    try:
                        shutil.rmtree(temp_config_dir)
                    except:
                        pass
                
                return None
            
            # Игнорируем некоторые предупреждения (как в dmed)
            if result.stderr:
                stderr_lower = result.stderr.lower()
                if 'javaldx: could not find a java runtime environment' in stderr_lower:
                    logger.debug("[PDF_CONV:JAVA_WARNING] Предупреждение о Java игнорируется (не критично)")
            
            # Очистка временной директории конфигурации для Windows
            if temp_config_dir and os.path.exists(temp_config_dir):
                try:
                    shutil.rmtree(temp_config_dir)
                except:
                    pass
            
            # Находим созданный PDF
            # LibreOffice создает PDF с тем же именем, но расширением .pdf
            docx_basename = os.path.splitext(os.path.basename(docx_path))[0]
            pdf_filename = docx_basename + '.pdf'
            pdf_path = os.path.join(output_dir, pdf_filename)
            
            log_pdf_conversion("PDF_SEARCH", "Поиск созданного PDF", expected_path=pdf_path)
            
            if not os.path.exists(pdf_path):
                logger.warning(f"[PDF_CONV:PDF_NOT_FOUND] Ожидаемый PDF не найден: {pdf_path}")
                # Пробуем найти любой PDF в директории
                pdf_files = [f for f in os.listdir(output_dir) if f.endswith('.pdf')]
                if pdf_files:
                    logger.info(f"[PDF_CONV:PDF_FOUND_ALT] Найдены альтернативные PDF: {pdf_files}")
                    pdf_path = os.path.join(output_dir, pdf_files[0])
                    log_pdf_conversion("PDF_USE_ALT", "Используется альтернативный PDF", pdf_path=pdf_path)
                else:
                    logger.error(f"[PDF_CONV:PDF_NOT_FOUND] PDF не найден в директории {output_dir}")
                    try:
                        dir_contents = os.listdir(output_dir)
                        logger.debug(f"[PDF_CONV:DIR_CONTENTS] Содержимое директории: {dir_contents[:20]}")
                    except Exception as e:
                        logger.error(f"[PDF_CONV:DIR_READ_ERROR] Ошибка при чтении директории: {e}")
                    return None
            
            pdf_size = os.path.getsize(pdf_path)
            log_pdf_conversion("PDF_FOUND", "PDF найден", pdf_path=pdf_path, pdf_size=pdf_size)
            
            # Сохраняем в MinIO
            log_pdf_conversion("STORAGE_SAVE", "Сохранение PDF в хранилище", 
                              pdf_size=pdf_size, uuid=document_uuid)
            with open(pdf_path, 'rb') as f:
                pdf_data = f.read()
            
            stored_path = storage_manager.save_file(
                pdf_data,
                f"{document_uuid}.pdf",
                'application/pdf'
            )
            
            log_pdf_conversion("STORAGE_SUCCESS", "PDF сохранен в хранилище", stored_path=stored_path)
            
            # Временная директория и файлы будут удалены автоматически при выходе из блока with
            
            log_pdf_conversion("SUCCESS", "Конвертация завершена успешно", 
                              stored_path=stored_path, pdf_size=pdf_size)
            return stored_path
        
    except Exception as e:
        log_error_with_context(e, f"convert_docx_to_pdf failed, docx_path={docx_path}, uuid={document_uuid}")
        return None


def find_libreoffice():
    """Находит путь к LibreOffice"""
    import platform
    
    if platform.system() == 'Windows':
        paths = [
            r'C:\Program Files\LibreOffice\program\soffice.exe',
            r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
        ]
    else:
        paths = [
            '/usr/bin/libreoffice',
            '/usr/bin/soffice',
            '/usr/local/bin/libreoffice',
            '/snap/bin/libreoffice',  # Snap пакет
        ]
    
    # Сначала проверяем стандартные пути
    for path in paths:
        if os.path.exists(path) and os.access(path, os.X_OK):
            logger.debug(f"[LIBREOFFICE:FOUND] Найден в стандартном пути: {path}")
            return path
    
    # Пробуем найти в PATH
    libreoffice_path = shutil.which('libreoffice')
    if libreoffice_path:
        logger.debug(f"[LIBREOFFICE:FOUND] Найден через PATH: {libreoffice_path}")
        return libreoffice_path
    
    soffice_path = shutil.which('soffice')
    if soffice_path:
        logger.debug(f"[LIBREOFFICE:FOUND] Найден soffice через PATH: {soffice_path}")
        return soffice_path
    
    logger.warning("[LIBREOFFICE:NOT_FOUND] LibreOffice не найден")
    return None

