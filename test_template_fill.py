#!/usr/bin/env python3
"""
Тест заполнения шаблона данными
"""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from app import create_app
from app.services.document import fill_docx_template, prepare_replacements
from docx import Document
import tempfile

app = create_app()

# Тестовые данные
test_data = {
    'uuid': 'test-uuid-12345',
    'mygov_doc_number': '227817784',
    'pin_code': '1234',
    'patient_name': 'ТЕСТОВЫЙ ПАЦИЕНТ',
    'gender': 'M',
    'age': '30',
    'jshshir': '12345678901234',
    'address': 'Тестовый адрес',
    'organization': 'Тестовая Организация',
    'diagnosis': 'Тестовый диагноз',
    'doctor_name': 'Тестовый Доктор',
    'doctor_position': 'Врач',
    'issue_date': '2025-11-30',
    'days_off_from': '2025-11-30',
    'days_off_to': '2025-12-05',
}

print("==========================================")
print("  Тест заполнения шаблона")
print("==========================================")

print("\n1. Подготовка замен:")
replacements = prepare_replacements(test_data)
print(f"Подготовлено {len(replacements)} замен:")
for key, value in list(replacements.items())[:10]:
    print(f"  {key} = '{value}'")

print("\n2. Загрузка шаблона...")
template_path = '/var/www/mygov-backend/templates/template_mygov.docx'
doc = Document(template_path)

print("\n3. Поиск плейсхолдеров в шаблоне до замены:")
import re
placeholder_pattern = r'\{\{[^}]+\}\}'
found_before = set()

for para in doc.paragraphs:
    placeholders = re.findall(placeholder_pattern, para.text)
    for ph in placeholders:
        found_before.add(ph)

print(f"Найдено плейсхолдеров: {len(found_before)}")
for ph in sorted(found_before):
    print(f"  - {ph}")

print("\n4. Заполнение шаблона...")
from app.services.document import replace_placeholders
replace_placeholders(doc, replacements)

print("\n5. Проверка после замены:")
found_after = set()
for para in doc.paragraphs:
    placeholders = re.findall(placeholder_pattern, para.text)
    for ph in placeholders:
        found_after.add(ph)

if found_after:
    print(f"⚠ Остались незамененные плейсхолдеры: {found_after}")
else:
    print("✓ Все плейсхолдеры заменены")

print("\n6. Проверка значений в документе:")
sample_text = ""
for para in doc.paragraphs[:5]:
    if para.text.strip():
        sample_text += para.text[:100] + "\n"

print("Пример текста из документа:")
print(sample_text)

print("\n7. Сохранение тестового файла...")
test_output = '/tmp/test_filled_template.docx'
doc.save(test_output)
print(f"✓ Сохранено: {test_output}")

print("\n==========================================")
print("  Тест завершен")
print("==========================================")

